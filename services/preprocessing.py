"""
Document Preprocessing Module

Handles text extraction from PDF and DOCX files, and provides text cleaning utilities.
"""

import re
import string
import pdfplumber
import docx
from bs4 import BeautifulSoup
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer, WordNetLemmatizer
from nltk.tokenize import word_tokenize, sent_tokenize

# Ensure NLTK data is available
import nltk
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')

class DocumentProcessor:
    """Encapsulates text extraction and cleaning functionalities."""

    def __init__(self, remove_stopwords=True, use_stemming=False, use_lemmatization=True):
        self.remove_stopwords = remove_stopwords
        self.use_stemming = use_stemming
        self.use_lemmatization = use_lemmatization
        self.stop_words = set(stopwords.words('english'))
        self.stemmer = PorterStemmer() if use_stemming else None
        self.lemmatizer = WordNetLemmatizer() if use_lemmatization else None

    def extract_text_from_pdf(self, file_path):
        """Extracts text from a PDF file."""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text

    def extract_text_from_docx(self, file_path):
        """Extracts text from a DOCX file."""
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    def clean_text(self, text):
        """Cleans raw text by removing noise and normalizing it."""
        # Remove HTML tags
        text = BeautifulSoup(text, "html.parser").get_text()
        
        # Remove URLs
        text = re.sub(r'https?://\S+|www\.\S+', '', text)
        
        # Remove punctuation
        text = text.translate(str.maketrans('', '', string.punctuation))
        
        # Convert to lowercase
        text = text.lower()
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Tokenize
        tokens = word_tokenize(text)
        
        # Process tokens
        processed_tokens = []
        for token in tokens:
            if self.remove_stopwords and token in self.stop_words:
                continue
            if self.use_stemming:
                token = self.stemmer.stem(token)
            if self.use_lemmatization:
                token = self.lemmatizer.lemmatize(token)
            processed_tokens.append(token)
            
        return " ".join(processed_tokens)

    def extract_text_from_file(self, file_path):
        """Auto-detects file type and extracts text accordingly."""
        if file_path.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path}")
    
    def preprocess_for_classification(self, text, max_length=512):
        """Preprocesses text specifically for document classification."""
        cleaned = self.clean_text(text)
        tokens = cleaned.split()
        if len(tokens) > max_length:
            tokens = tokens[:max_length]
        return " ".join(tokens)
    
    def preprocess_for_summarization(self, text):
        """Preprocesses text specifically for summarization (lighter cleaning)."""
        # Remove HTML tags
        text = BeautifulSoup(text, "html.parser").get_text()
        # Remove URLs
        text = re.sub(r'https?://\S+|www\.\S+', '', text)
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    
    def split_into_sentences(self, text):
        """Splits text into sentences for processing."""
        return sent_tokenize(text)
    
    def extract_metadata(self, file_path):
        """Extracts basic metadata from documents."""
        metadata = {'file_path': file_path, 'file_type': None}
        
        if file_path.lower().endswith('.pdf'):
            metadata['file_type'] = 'pdf'
            with pdfplumber.open(file_path) as pdf:
                metadata['num_pages'] = len(pdf.pages)
                if pdf.metadata:
                    metadata.update(pdf.metadata)
        elif file_path.lower().endswith('.docx'):
            metadata['file_type'] = 'docx'
            doc = docx.Document(file_path)
            metadata['num_paragraphs'] = len(doc.paragraphs)
            if doc.core_properties:
                props = doc.core_properties
                metadata.update({
                    'title': props.title,
                    'author': props.author,
                    'created': props.created,
                    'modified': props.modified
                })
        
        return metadata

